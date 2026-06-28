#!/usr/bin/env python3
"""
Round-trip quality evaluator for SDS Converter.

Compares original JSON (output from PDF via LLM) with text extracted from
DOCX generated from that JSON, evaluating field preservation fidelity.

Usage: python3 compare_json.py
"""

import json
import os
import sys
from docx import Document

SAMPLES_DIR = "/Users/k_tanabe/Documents/Documents/oss_rust/sds_converter/samples"
RT_OUTPUT_DIR = os.path.join(SAMPLES_DIR, "rt_output")

# All top-level section keys in MHLW SDS JSON
SDS_SECTIONS = [
    "Identification",
    "Composition",
    "HazardIdentification",
    "FirstAidMeasures",
    "FireFightingMeasures",
    "AccidentalReleaseMeasures",
    "HandlingAndStorage",
    "ExposureControlPersonalProtection",
    "PhysicalChemicalProperties",
    "StabilityReactivity",
    "ToxicologicalInformation",
    "EcologicalInformation",
    "DisposalConsiderations",
    "TransportInformation",
    "RegulatoryInformation",
    "OtherInformation",
    "Datasheet",
]


def extract_all_text_values(obj, prefix=""):
    """Recursively extract all string values from a JSON object."""
    texts = []
    if isinstance(obj, dict):
        for k, v in obj.items():
            child_prefix = f"{prefix}.{k}" if prefix else k
            texts.extend(extract_all_text_values(v, child_prefix))
    elif isinstance(obj, list):
        for i, item in enumerate(obj):
            child_prefix = f"{prefix}[{i}]"
            texts.extend(extract_all_text_values(item, child_prefix))
    elif isinstance(obj, str) and obj.strip():
        texts.append((prefix, obj))
    elif isinstance(obj, (int, float)) and obj != 0:
        texts.append((prefix, str(obj)))
    return texts


def extract_full_texts(obj, prefix=""):
    """Extract only FullText fields from JSON."""
    texts = []
    if isinstance(obj, dict):
        for k, v in obj.items():
            child_prefix = f"{prefix}.{k}" if prefix else k
            if k == "FullText" and isinstance(v, str) and v.strip():
                texts.append((child_prefix, v))
            else:
                texts.extend(extract_full_texts(v, child_prefix))
    elif isinstance(obj, list):
        for i, item in enumerate(obj):
            texts.extend(extract_full_texts(item, f"{prefix}[{i}]"))
    return texts


def extract_key_fields(data):
    """Extract key identifying fields from JSON."""
    fields = {}
    # Product name / identification
    ident = data.get("Identification", {})
    product = ident.get("SDSProductName", {})
    if isinstance(product, dict):
        fields["product_name"] = product.get("ProductName", "")
    elif isinstance(product, str):
        fields["product_name"] = product

    # Issue date
    datasheet = data.get("Datasheet", {})
    fields["issue_date"] = datasheet.get("IssueDate", "")
    fields["revision_date"] = datasheet.get("RevisionDate", "")

    # Supplier/manufacturer
    supplier = ident.get("SupplierManufacturer", {})
    if isinstance(supplier, list) and supplier:
        supplier = supplier[0]
    if isinstance(supplier, dict):
        fields["company"] = supplier.get("CompanyName", "")

    # Hazard classification
    hazard = data.get("HazardIdentification", {})
    if isinstance(hazard, dict):
        ghs = hazard.get("GHSHazardClassification", {})
        if isinstance(ghs, dict):
            fields["ghs_categories"] = list(ghs.keys())

    # Composition
    comp = data.get("Composition", {})
    if isinstance(comp, dict):
        comps = comp.get("CompositionAndConcentration", [])
        substances = []
        for c in comps:
            if isinstance(c, dict):
                si = c.get("SubstanceIdentifiers", {})
                names = si.get("SubstanceNames", {})
                name = names.get("GenericName", names.get("SDSName", ""))
                if name:
                    substances.append(name)
        fields["substances"] = substances

    return fields


def get_docx_text(docx_path):
    """Extract all text from a DOCX file."""
    try:
        doc = Document(docx_path)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        paragraphs.append(cell_text)
        return "\n".join(paragraphs)
    except Exception as e:
        return f"ERROR: {e}"


def check_text_in_docx(text_value, docx_text):
    """Check how much of a text value is preserved in the DOCX."""
    if not text_value or not text_value.strip():
        return None  # Skip empty
    # Split into meaningful chunks (sentences or significant phrases)
    # For Japanese text, use character-level match
    text_clean = text_value.strip()
    if len(text_clean) < 5:
        return None  # Too short to be meaningful
    # Check if the text appears (substring match)
    if text_clean in docx_text:
        return 1.0
    # Partial match: count overlapping characters in windows
    # Use a sliding window approach
    window = 10
    matches = 0
    total_windows = 0
    for i in range(0, len(text_clean) - window + 1, window // 2):
        chunk = text_clean[i:i+window]
        total_windows += 1
        if chunk in docx_text:
            matches += 1
    if total_windows == 0:
        return 1.0 if text_clean in docx_text else 0.0
    return matches / total_windows


def evaluate_file(file_num):
    """Evaluate round-trip quality for a single file."""
    json_path = os.path.join(SAMPLES_DIR, f"output{file_num}.json")
    docx_path = os.path.join(RT_OUTPUT_DIR, f"rt_{file_num}.docx")

    if not os.path.exists(json_path):
        return None, f"JSON not found: {json_path}"
    if not os.path.exists(docx_path):
        return None, f"DOCX not found: {docx_path}"

    with open(json_path, encoding="utf-8") as f:
        data = json.load(f)

    docx_text = get_docx_text(docx_path)
    if docx_text.startswith("ERROR:"):
        return None, docx_text

    # Extract all string values from JSON
    all_values = extract_all_text_values(data)
    full_texts = extract_full_texts(data)
    key_fields = extract_key_fields(data)

    # Sections present in JSON
    sections_in_json = [s for s in SDS_SECTIONS if s in data]
    sections_missing = [s for s in SDS_SECTIONS if s not in data]

    # Evaluate field preservation
    scores = []
    section_scores = {}

    for path, value in all_values:
        score = check_text_in_docx(value, docx_text)
        if score is not None:
            scores.append((path, value[:50], score))
            # Extract section name (first component)
            section = path.split(".")[0]
            if section not in section_scores:
                section_scores[section] = []
            section_scores[section].append(score)

    # Overall stats
    if scores:
        overall_retention = sum(s for _, _, s in scores) / len(scores)
    else:
        overall_retention = 0.0

    # Per-section retention
    section_retention = {}
    for sec, sec_scores in section_scores.items():
        section_retention[sec] = sum(sec_scores) / len(sec_scores)

    # FullText-specific retention
    ft_scores = []
    for path, ft_value in full_texts:
        score = check_text_in_docx(ft_value, docx_text)
        if score is not None:
            ft_scores.append(score)
    fulltext_retention = sum(ft_scores) / len(ft_scores) if ft_scores else None

    result = {
        "file_num": file_num,
        "json_path": json_path,
        "docx_path": docx_path,
        "sections_in_json": sections_in_json,
        "sections_missing_from_json": sections_missing,
        "total_fields_checked": len(scores),
        "overall_retention": overall_retention,
        "fulltext_retention": fulltext_retention,
        "section_retention": section_retention,
        "key_fields": key_fields,
        "low_score_fields": [(p, v, s) for p, v, s in scores if s < 0.3],
        "full_texts_count": len(full_texts),
    }
    return result, None


def main():
    file_nums = ["02", "03", "04", "05", "06", "07", "08", "09", "10", "11", "12"]
    results = []
    errors = []

    print("=" * 70)
    print("SDS Converter Round-Trip Quality Evaluation")
    print("=" * 70)

    for fn in file_nums:
        result, error = evaluate_file(fn)
        if error:
            print(f"[{fn}] ERROR: {error}")
            errors.append((fn, error))
        else:
            results.append(result)
            retention_pct = result["overall_retention"] * 100
            ft_pct = (result["fulltext_retention"] or 0) * 100
            print(
                f"[{fn}] 総フィールド保持率: {retention_pct:.1f}%  "
                f"FullText保持率: {ft_pct:.1f}%  "
                f"セクション数: {len(result['sections_in_json'])}/17"
            )

    print()
    print("=" * 70)
    print("Per-Section Retention Rates (averaged across all files)")
    print("=" * 70)
    # Aggregate section scores
    all_section_scores = {}
    for r in results:
        for sec, score in r["section_retention"].items():
            if sec not in all_section_scores:
                all_section_scores[sec] = []
            all_section_scores[sec].append(score)

    section_avg = {sec: sum(scores) / len(scores) for sec, scores in all_section_scores.items()}
    for sec, avg in sorted(section_avg.items(), key=lambda x: -x[1]):
        print(f"  {sec:<45} {avg*100:.1f}%")

    # Best / worst files
    if results:
        best = max(results, key=lambda r: r["overall_retention"])
        worst = min(results, key=lambda r: r["overall_retention"])
        print()
        print(f"最良ファイル: output{best['file_num']}.json ({best['overall_retention']*100:.1f}%)")
        print(f"最悪ファイル: output{worst['file_num']}.json ({worst['overall_retention']*100:.1f}%)")

    return results, errors, section_avg


if __name__ == "__main__":
    results, errors, section_avg = main()

    # Save detailed results as JSON for the report
    output = {
        "results": [
            {k: v for k, v in r.items() if k != "low_score_fields"}
            for r in results
        ],
        "errors": errors,
        "section_averages": {k: v for k, v in section_avg.items()},
    }
    out_path = os.path.join(RT_OUTPUT_DIR, "evaluation_results.json")
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(output, f, ensure_ascii=False, indent=2)
    print(f"\n詳細結果を保存: {out_path}")
